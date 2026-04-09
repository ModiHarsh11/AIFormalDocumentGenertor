"""Views for Office Order document generation, preview, and export."""

import logging

from django.conf import settings
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.template.loader import render_to_string
from django.utils import timezone
from django.views.decorators.http import require_POST

from generator.views.common import save_document_to_db
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from weasyprint import HTML, CSS

from generator.constants import DESIGNATION_MAP

from .common import (
    OFFICE_ORDER, format_date_ddmmyyyy,
    handle_generate_body, handle_regenerate_body, handle_update_body,
    make_docx_response,
)

logger = logging.getLogger(__name__)
_ORG_SUFFIX = "BISAG-N"
_OFFICE_LINE_FONT = "Nirmala UI"


def _format_office_recipient(designation: str) -> str:
    """Append organization suffix for Office Order recipients."""
    text = (designation or "").strip()
    if not text:
        return text
    if _ORG_SUFFIX.lower() in text.lower():
        return text
    return f"{text}, {_ORG_SUFFIX}"


def _set_run_font(run, font_name: str) -> None:
    """Force a single font for mixed-script runs (Latin + Devanagari)."""
    run.font.name = font_name
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


# ---------------------------------------------------------------------------
# AI body generation
# ---------------------------------------------------------------------------

@require_POST
@login_required
def generate_body(request):
    """Generate an Office Order body via the AI service."""
    return handle_generate_body(request, "office", "generate_body")


@require_POST
@login_required
def regenerate_office_body(request):
    """Regenerate an Office Order body with refinement feedback."""
    return handle_regenerate_body(request, "office", "regenerate_office_body")


# ---------------------------------------------------------------------------
# Form & preview
# ---------------------------------------------------------------------------

def office_order_form(request):
    return render(request, "generator/office_order_form.html", {
        "designations": DESIGNATION_MAP.keys(),
    })


@login_required
def result_office_order(request):
    if request.method != "POST":
        return redirect("home")

    lang = request.POST.get("language", "en")
    raw_date = request.POST.get("date") or timezone.now().strftime("%Y-%m-%d")
    date = format_date_ddmmyyyy(raw_date)

    reference = request.POST.get("reference", "").strip()
    if not reference:
        reference = (
            "बायसेग-एन/कार्यालय आदेश/2026/"
            if lang == "hi"
            else "BISAG-N/Office Order/2026/"
        )

    body_prompt = request.POST.get("body_prompt", "").strip()
    from_position = request.POST.get("from_position", "")
    to_recipients = request.POST.getlist("to_recipients[]")

    data = {
        "language": lang,
        "header": OFFICE_ORDER["header"][lang],
        "title": OFFICE_ORDER["title_hi"] if lang == "hi" else OFFICE_ORDER["title_en"],
        "reference": reference,
        "date": date,
        "date_raw": raw_date,
        "body": request.POST.get("body", "").strip(),
        "body_prompt": body_prompt,
        "from": DESIGNATION_MAP.get(from_position, {}).get(lang, ""),
        "from_position": from_position,
        "to": [
            _format_office_recipient(DESIGNATION_MAP.get(x, {}).get(lang, x))
            for x in to_recipients
        ],
        "to_data": to_recipients,
    }

    # Persist to database
    doc = save_document_to_db(
        user=request.user,
        document_type="office",
        language=lang,
        reference_id=reference,
        date_raw=raw_date,
        body_prompt=body_prompt,
        body=data["body"],
        from_position=from_position,
        to_data=to_recipients,
    )
    data["serial_number"] = doc.serial_number
    data["doc_id"] = doc.id

    request.session["doc_data"] = data
    return render(request, "generator/result_office_order.html", data)


@require_POST
@login_required
def update_office_body(request):
    return handle_update_body(request, "doc_data")


@login_required
def office_order_form(request):
    return render(request, "generator/office_order_form.html", {
        "designations": DESIGNATION_MAP.keys(),
    })


# ---------------------------------------------------------------------------
# PDF & DOCX export
# ---------------------------------------------------------------------------

def download_pdf(request):
    data = request.session.get("doc_data")
    if not data:
        return HttpResponse("No office order generated", status=400)

    html = render_to_string("generator/pdf_office_order.html", data)

    pdf = HTML(
        string=html,
        base_url=str(settings.BASE_DIR),
    ).write_pdf(
        stylesheets=[
            CSS(string="""
                @page { size: A4; margin: 2.5cm; }
                body { font-family: 'Times New Roman', serif; font-size: 12pt; line-height: 1.6; }
                .center { text-align: center; }
                .bold { font-weight: bold; }
                .ref-date-row { display: table; width: 100%; margin: 20px 0; }
                .ref-left { display: table-cell; text-align: left; font-weight: bold; width: 50%; font-family: 'Nirmala UI', 'Times New Roman', serif; }
                .date-right { display: table-cell; text-align: right; font-weight: bold; width: 50%; font-family: 'Nirmala UI', 'Times New Roman', serif; }
                .title { text-align: center; font-weight: bold; text-decoration: underline; margin: 20px 0; font-size: 14pt; }
                .body { text-align: justify; margin: 20px 0; white-space: pre-wrap; }
                .from-section { text-align: right; font-weight: bold; margin: 40px 0 20px; }
                .to-section { margin-top: 20px; }
                .to-section div { margin: 5px 0; font-weight: bold; }
            """)
        ],
    )

    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="Office_Order.pdf"'
    return response


def download_docx(request):
    data = request.session.get("doc_data")
    if not data:
        return HttpResponse("No office order generated", status=400)

    doc = Document()

    # Header
    for line in data["header"]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

    # Reference & Date
    p = doc.add_paragraph(f"Ref: {data['reference']}")
    p.runs[0].bold = True
    _set_run_font(p.runs[0], _OFFICE_LINE_FONT)

    p = doc.add_paragraph(f"Date: {data['date']}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True
    _set_run_font(p.runs[0], _OFFICE_LINE_FONT)

    # Title
    p = doc.add_paragraph(data["title"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].underline = True

    # Body
    doc.add_paragraph(data["body"])

    # From
    p = doc.add_paragraph(data["from"])
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True

    # To
    for t in data["to"]:
        p = doc.add_paragraph(t)
        p.runs[0].bold = True

    return make_docx_response(doc, "Office_Order.docx")

