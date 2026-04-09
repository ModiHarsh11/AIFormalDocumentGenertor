"""Views for Circular document generation, preview, and export."""

import logging

from django.conf import settings
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.template.loader import render_to_string
from django.utils import timezone
from django.views.decorators.http import require_POST
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from weasyprint import HTML

from generator.constants import DESIGNATION_MAP

from .common import (
    CIRCULAR, format_date_ddmmyyyy,
    handle_generate_body, handle_regenerate_body, handle_update_body,
    maybe_translate_subject,
    get_logo_path, get_logo_file_uri, make_docx_response,
    save_document_to_db,
)

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# AI body generation
# ---------------------------------------------------------------------------

@require_POST
@login_required
def generate_circular_body(request):
    """Generate a Circular body via the AI service."""
    return handle_generate_body(request, "circular", "generate_circular_body")


@require_POST
@login_required
def regenerate_circular_body(request):
    """Regenerate a Circular body with refinement feedback."""
    return handle_regenerate_body(request, "circular", "regenerate_circular_body")


# ---------------------------------------------------------------------------
# Form & preview
# ---------------------------------------------------------------------------

def circular_form(request):
    return render(request, "generator/circular_form.html", {
        "people": CIRCULAR["people"],
    })


@login_required
def result_circular(request):
    if request.method != "POST":
        return redirect("circular_form")

    lang = request.POST.get("language") or "en"
    raw_date = request.POST.get("date") or timezone.now().strftime("%Y-%m-%d")
    date = (
        format_date_ddmmyyyy(raw_date).replace("-", "/")
        if raw_date
        else timezone.now().strftime("%d/%m/%Y")
    )
    subject = maybe_translate_subject(request.POST.get("subject"), lang)
    body = request.POST.get("body")
    body_prompt = request.POST.get("body_prompt", "").strip()

    from_position = request.POST.get("from_position")
    from_designation = DESIGNATION_MAP.get(from_position, {}).get(lang, "") if from_position else ""

    to_ids = request.POST.getlist("to[]")
    people = CIRCULAR["people"]
    to_people = [p for p in people if str(p["id"]) in to_ids]

    lang_key = "hindi" if lang == "hi" else "english"
    header = {
        "org_name": CIRCULAR["header"][lang_key]["org_name"],
        "ministry": CIRCULAR["header"][lang_key]["ministry"],
        "government": CIRCULAR["header"][lang_key]["government"],
    }

    data = {
        "language": lang,
        "header": header,
        "date": date,
        "date_raw": raw_date,
        "subject": subject,
        "body": body,
        "body_prompt": body_prompt,
        "from": from_designation,
        "from_position": from_position or "",
        "from_org": "BISAG-N",
        "to_people": to_people,
        "to_data": to_ids,
    }

    # Persist to database
    doc = save_document_to_db(
        user=request.user,
        document_type="circular",
        language=lang,
        date_raw=raw_date,
        subject=subject,
        body_prompt=body_prompt,
        body=body or "",
        from_position=from_position or "",
        to_data=to_ids,
    )
    data["serial_number"] = doc.serial_number
    data["doc_id"] = doc.id

    request.session["circular_data"] = data
    return render(request, "generator/result_circular.html", data)


@require_POST
@login_required
def update_circular_body(request):
    return handle_update_body(request, "circular_data")


# ---------------------------------------------------------------------------
# PDF & DOCX export
# ---------------------------------------------------------------------------

def download_circular_pdf(request):
    data = request.session.get("circular_data")
    if not data:
        return HttpResponse("No circular generated", status=400)

    # Work on a copy to avoid polluting the session with logo_path
    render_data = {**data}
    logo_uri = get_logo_file_uri()
    if logo_uri:
        render_data["logo_path"] = logo_uri

    html = render_to_string("generator/pdf_circular.html", render_data)

    pdf = HTML(
        string=html,
        base_url=str(settings.BASE_DIR),
    ).write_pdf(
        optimize_images=True,
        jpeg_quality=85,
        presentational_hints=True,
    )

    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="Circular.pdf"'
    return response


def download_circular_docx(request):
    data = request.session.get("circular_data")
    if not data:
        return HttpResponse("No circular generated", status=400)

    doc = Document()
    lang = data.get("language", "en")

    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add BISAG Logo
    logo_path = get_logo_path()
    if logo_path:
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, height=Inches(0.9))
        doc.add_paragraph()

    # Header lines
    for line in data["header"].values():
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(14)

    # Circular title
    title_text = "परिपत्र" if lang == "hi" else "Circular"
    p = doc.add_paragraph(title_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].underline = True
    p.runs[0].font.size = Pt(16)

    # Date
    date_label = "दिनांक :" if lang == "hi" else "Date :"
    p = doc.add_paragraph(f"{date_label} {data['date']}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    # Subject
    subject_label = "विषय :" if lang == "hi" else "Subject :"
    p = doc.add_paragraph(f"{subject_label} {data['subject']}")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    # Body
    p = doc.add_paragraph(data["body"])
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].font.size = Pt(12)

    # From section
    p = doc.add_paragraph(f"{data['from']}\n{data.get('from_org', 'BISAG-N')}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(12)

    # Space before table
    doc.add_paragraph()

    # To section — Table
    to_people = data.get("to_people", [])
    if to_people:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        sr_label = "क्र." if lang == "hi" else "S. No."
        name_label = "नाम" if lang == "hi" else "Name"
        sign_label = "हस्ताक्षर" if lang == "hi" else "Sign"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = sr_label
        hdr_cells[1].text = name_label
        hdr_cells[2].text = sign_label

        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, person in enumerate(to_people, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = f"{idx}."
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            name = person.get("name_hi") if lang == "hi" else person.get("name_en")
            designation = person.get("designation_hi") if lang == "hi" else person.get("designation_en")
            if name and designation:
                row_cells[1].text = f"{name} ({designation})"
            else:
                row_cells[1].text = name or designation or ""
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            row_cells[2].text = ""
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.columns[0].width = Inches(1.0)
        table.columns[1].width = Inches(3.5)
        table.columns[2].width = Inches(1.5)

    return make_docx_response(doc, "Circular.docx")

