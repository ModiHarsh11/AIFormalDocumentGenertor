"""Dashboard, document history, view by ID, edit, and download-by-ID views."""

import base64
import logging
import os
import uuid

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator
from django.http import HttpResponse, Http404
from django.shortcuts import render, redirect, get_object_or_404
from django.template.loader import render_to_string
from django.utils import timezone
from django.views.decorators.http import require_POST
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from io import BytesIO
from weasyprint import HTML, CSS

from generator.constants import DESIGNATION_MAP
from generator.models import GeneratedDocument

from .common import (
    OFFICE_ORDER, CIRCULAR, POLICY,
    format_date_ddmmyyyy, format_date_ddmmmyyyy,
    get_logo_file_uri, make_docx_response,
)
from django.conf import settings

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Helpers: Reconstruct session-compatible data dict from GeneratedDocument
# ---------------------------------------------------------------------------

def _office_recipients(positions, lang):
    """Convert a list of position keys to formatted recipient strings."""
    out = []
    for pos in positions:
        text = DESIGNATION_MAP.get(pos, {}).get(lang, pos)
        suffix = "BISAG-N"
        if suffix.lower() not in text.lower():
            text = f"{text}, {suffix}"
        out.append(text)
    return out


def _reconstruct_data(doc: GeneratedDocument) -> dict:
    """Build the same dict structure the result views produce, from a DB record."""
    lang = doc.language
    if doc.document_type == "office":
        header = OFFICE_ORDER["header"][lang]
        return {
            "language": lang,
            "header": header,
            "title": OFFICE_ORDER["title_hi"] if lang == "hi" else OFFICE_ORDER["title_en"],
            "reference": doc.reference_id,
            "date": format_date_ddmmyyyy(doc.date_raw) if doc.date_raw else "",
            "date_raw": doc.date_raw,
            "body": doc.body,
            "body_prompt": doc.body_prompt,
            "from": DESIGNATION_MAP.get(doc.from_position, {}).get(lang, doc.from_position),
            "from_position": doc.from_position,
            "to": _office_recipients(doc.to_data or [], lang),
            "to_data": doc.to_data or [],
            "serial_number": doc.serial_number,
            "doc_id": doc.id,
            "version": doc.version,
        }

    elif doc.document_type == "circular":
        lang_key = "hindi" if lang == "hi" else "english"
        header = {
            "org_name": CIRCULAR["header"][lang_key]["org_name"],
            "ministry": CIRCULAR["header"][lang_key]["ministry"],
            "government": CIRCULAR["header"][lang_key]["government"],
        }
        people = CIRCULAR["people"]
        to_ids = [str(x) for x in (doc.to_data or [])]
        to_people = [p for p in people if str(p["id"]) in to_ids]
        raw_date = doc.date_raw
        date_fmt = (format_date_ddmmyyyy(raw_date).replace("-", "/") if raw_date
                    else timezone.now().strftime("%d/%m/%Y"))
        return {
            "language": lang,
            "header": header,
            "date": date_fmt,
            "date_raw": doc.date_raw,
            "subject": doc.subject,
            "body": doc.body,
            "body_prompt": doc.body_prompt,
            "from": DESIGNATION_MAP.get(doc.from_position, {}).get(lang, doc.from_position),
            "from_position": doc.from_position,
            "from_org": "BISAG-N",
            "to_people": to_people,
            "to_data": doc.to_data or [],
            "serial_number": doc.serial_number,
            "doc_id": doc.id,
            "version": doc.version,
        }

    elif doc.document_type == "policy":
        lang_key = "hi" if lang == "hi" else "en"
        header = {
            "org_name": POLICY["header"][lang_key][0],
            "ministry": POLICY["header"][lang_key][1],
            "government": POLICY["header"][lang_key][2],
        }
        to_designations = [DESIGNATION_MAP.get(x, {}).get(lang, x) for x in (doc.to_data or [])]
        date_fmt = format_date_ddmmmyyyy(doc.date_raw) if doc.date_raw else ""
        extra = doc.extra_data or {}
        pdf_path = extra.get("uploaded_pdf_path", "")
        # Regenerate base64 preview if file still exists
        pdf_base64 = None
        if pdf_path and os.path.exists(pdf_path):
            try:
                with open(pdf_path, "rb") as f:
                    pdf_base64 = base64.b64encode(f.read()).decode("utf-8")
            except Exception:
                pass
        return {
            "language": lang,
            "header": header,
            "date": date_fmt,
            "date_raw": doc.date_raw,
            "subject": doc.subject,
            "body": doc.body,
            "body_prompt": doc.body_prompt,
            "from": DESIGNATION_MAP.get(doc.from_position, {}).get(lang, doc.from_position),
            "from_position": doc.from_position,
            "to_designations": to_designations,
            "to_data": doc.to_data or [],
            "attached_pdf_name": extra.get("attached_pdf_name", ""),
            "uploaded_pdf_path": pdf_path,
            "pdf_base64": pdf_base64,
            "serial_number": doc.serial_number,
            "doc_id": doc.id,
            "version": doc.version,
        }

    return {}


# ---------------------------------------------------------------------------
# Dashboard
# ---------------------------------------------------------------------------

@login_required
def dashboard(request):
    """User's document history dashboard."""
    doc_type_filter = request.GET.get("type", "")
    search_query    = request.GET.get("q", "").strip()

    queryset = GeneratedDocument.objects.filter(
        user=request.user,
        is_deleted=False,
    ).order_by("-created_at")

    if doc_type_filter:
        queryset = queryset.filter(document_type=doc_type_filter)
    if search_query:
        from django.db.models import Q
        queryset = queryset.filter(
            Q(serial_number__icontains=search_query)
            | Q(subject__icontains=search_query)
            | Q(reference_id__icontains=search_query)
        )

    paginator = Paginator(queryset, 15)
    page_obj  = paginator.get_page(request.GET.get("page"))

    stats = {
        "total": GeneratedDocument.objects.filter(user=request.user, is_deleted=False).count(),
        "office": GeneratedDocument.objects.filter(user=request.user, document_type="office", is_deleted=False).count(),
        "circular": GeneratedDocument.objects.filter(user=request.user, document_type="circular", is_deleted=False).count(),
        "policy": GeneratedDocument.objects.filter(user=request.user, document_type="policy", is_deleted=False).count(),
        "downloads": sum(
            d.download_count for d in
            GeneratedDocument.objects.filter(user=request.user, is_deleted=False).only("download_count")
        ),
    }

    return render(request, "generator/dashboard.html", {
        "page_obj": page_obj,
        "stats": stats,
        "doc_type_filter": doc_type_filter,
        "search_query": search_query,
    })


# ---------------------------------------------------------------------------
# View document by ID (reconstruct result from DB)
# ---------------------------------------------------------------------------

@login_required
def document_view(request, doc_id):
    """Re-render a previously generated document from its DB record."""
    doc = get_object_or_404(
        GeneratedDocument,
        id=doc_id,
        is_deleted=False,
    )
    if doc.user != request.user and not request.user.is_staff:
        raise Http404

    data = _reconstruct_data(doc)
    # Update session so existing download views still work
    session_key = _session_key(doc.document_type)
    request.session[session_key] = data

    template_map = {
        "office":   "generator/result_office_order.html",
        "circular": "generator/result_circular.html",
        "policy":   "generator/result_policy.html",
    }
    template = template_map.get(doc.document_type)
    if not template:
        raise Http404

    return render(request, template, data)


def _session_key(doc_type: str) -> str:
    return {"office": "doc_data", "circular": "circular_data", "policy": "policy_data"}.get(doc_type, "doc_data")


# ---------------------------------------------------------------------------
# Edit document fields
# ---------------------------------------------------------------------------

@login_required
def document_edit(request, doc_id):
    """Edit individual fields of a previously generated document."""
    doc = get_object_or_404(GeneratedDocument, id=doc_id, is_deleted=False)
    if doc.user != request.user and not request.user.is_staff:
        raise Http404

    if request.method == "POST":
        doc.language      = request.POST.get("language", doc.language)
        doc.reference_id  = request.POST.get("reference_id", "").strip()
        doc.date_raw      = request.POST.get("date_raw", "").strip()
        doc.subject       = request.POST.get("subject", "").strip()
        doc.body          = request.POST.get("body", "").strip()
        doc.from_position = request.POST.get("from_position", "").strip()

        raw_to = request.POST.getlist("to_data[]")
        doc.to_data = [x.strip() for x in raw_to if x.strip()]

        # Handle policy PDF attachment
        if doc.document_type == "policy":
            extra = dict(doc.extra_data or {})
            new_pdf = request.FILES.get("policy_pdf")
            if new_pdf:
                # Save new PDF, replace old one
                upload_dir = os.path.join(settings.MEDIA_ROOT, "policy_uploads")
                os.makedirs(upload_dir, exist_ok=True)
                pdf_filename = f"policy_{uuid.uuid4().hex}.pdf"
                new_pdf_path = os.path.join(upload_dir, pdf_filename)
                with open(new_pdf_path, "wb+") as dest:
                    for chunk in new_pdf.chunks():
                        dest.write(chunk)
                # Remove old file if it exists
                old_path = extra.get("uploaded_pdf_path", "")
                if old_path and os.path.exists(old_path):
                    try:
                        os.remove(old_path)
                    except Exception:
                        pass
                extra["uploaded_pdf_path"] = new_pdf_path
                extra["attached_pdf_name"] = new_pdf.name
            # else: keep existing uploaded_pdf_path
            doc.extra_data = extra

        doc.version += 1
        doc.save()

        messages.success(request, f"Document {doc.serial_number} updated (version {doc.version}).")
        return redirect("document_view", doc_id=doc.id)

    # Prepare dropdown options based on doc type
    context = {
        "doc": doc,
        "designations": list(DESIGNATION_MAP.keys()),
    }
    if doc.document_type == "circular":
        context["people"] = CIRCULAR["people"]
    elif doc.document_type == "office":
        context["office_to_options"] = list(DESIGNATION_MAP.keys())
    elif doc.document_type == "policy":
        context["policy_to_options"] = list(DESIGNATION_MAP.keys())
        context["current_pdf_name"] = (doc.extra_data or {}).get("attached_pdf_name", "")
        context["has_pdf"] = bool((doc.extra_data or {}).get("uploaded_pdf_path", ""))

    return render(request, "generator/document_edit.html", context)


# ---------------------------------------------------------------------------
# Soft-delete document
# ---------------------------------------------------------------------------

@login_required
@require_POST
def document_delete(request, doc_id):
    doc = get_object_or_404(GeneratedDocument, id=doc_id, is_deleted=False)
    if doc.user != request.user and not request.user.is_staff:
        raise Http404
    doc.is_deleted = True
    doc.save(update_fields=["is_deleted"])
    messages.success(request, f"Document {doc.serial_number} deleted.")
    return redirect("dashboard")


# ---------------------------------------------------------------------------
# Download PDF by doc_id (from history)
# ---------------------------------------------------------------------------

@login_required
def download_doc_pdf(request, doc_id):
    doc = get_object_or_404(GeneratedDocument, id=doc_id, is_deleted=False)
    if doc.user != request.user and not request.user.is_staff:
        raise Http404

    doc.download_count += 1
    doc.save(update_fields=["download_count"])

    data = _reconstruct_data(doc)
    logo_uri = get_logo_file_uri()
    if logo_uri:
        data["logo_path"] = logo_uri

    template_map = {
        "office":   "generator/pdf_office_order.html",
        "circular": "generator/pdf_circular.html",
        "policy":   "generator/pdf_policy.html",
    }
    template = template_map.get(doc.document_type)
    if not template:
        raise Http404

    html_string = render_to_string(template, data)
    first_page_pdf = HTML(string=html_string, base_url=".").write_pdf()

    # For policy: merge with the attached PDF if it exists
    final_pdf = first_page_pdf
    if doc.document_type == "policy":
        pdf_path = (doc.extra_data or {}).get("uploaded_pdf_path", "")
        if pdf_path and os.path.exists(pdf_path):
            try:
                from pypdf import PdfReader, PdfWriter
                writer = PdfWriter()
                for page in PdfReader(BytesIO(first_page_pdf)).pages:
                    writer.add_page(page)
                for page in PdfReader(pdf_path).pages:
                    writer.add_page(page)
                buf = BytesIO()
                writer.write(buf)
                final_pdf = buf.getvalue()
            except Exception as e:
                logger.error("PDF merge error in download_doc_pdf: %s", e)

    fname_map = {
        "office": "Office_Order",
        "circular": "Circular",
        "policy": "Policy",
    }
    fname = f"{fname_map[doc.document_type]}_{doc.serial_number}.pdf"
    response = HttpResponse(final_pdf, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{fname}"'
    return response


# ---------------------------------------------------------------------------
# Download DOCX by doc_id (from history) — lean re-creation
# ---------------------------------------------------------------------------

@login_required
def download_doc_docx(request, doc_id):
    doc = get_object_or_404(GeneratedDocument, id=doc_id, is_deleted=False)
    if doc.user != request.user and not request.user.is_staff:
        raise Http404

    doc.download_count += 1
    doc.save(update_fields=["download_count"])

    data = _reconstruct_data(doc)
    word_doc = Document()

    # Header
    header_lines = data.get("header", []) or list(data.get("header", {}).values())
    if isinstance(header_lines, dict):
        header_lines = list(header_lines.values())
    for line in header_lines:
        p = word_doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

    # Reference & Date
    ref = data.get("reference", data.get("reference_id", ""))
    if ref:
        p = word_doc.add_paragraph(f"Ref: {ref}")
        p.runs[0].bold = True

    date_val = data.get("date", "")
    if date_val:
        p = word_doc.add_paragraph(f"Date: {date_val}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].bold = True

    # Subject
    subject = data.get("subject", "")
    if subject:
        p = word_doc.add_paragraph(f"Subject: {subject}")
        p.runs[0].bold = True

    # Title (Office Order)
    title = data.get("title", "")
    if title:
        p = word_doc.add_paragraph(title)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

    # Body
    body = data.get("body", "")
    if body:
        word_doc.add_paragraph(body)

    # From
    from_val = data.get("from", "")
    if from_val:
        p = word_doc.add_paragraph(from_val)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].bold = True

    # To
    for t in data.get("to", []):
        p = word_doc.add_paragraph(t)
        p.runs[0].bold = True
    for t in data.get("to_designations", []):
        p = word_doc.add_paragraph(t)
        p.runs[0].bold = True
    for p_entry in data.get("to_people", []):
        p = word_doc.add_paragraph(p_entry.get("name", str(p_entry)))
        p.runs[0].bold = True

    fname_map = {"office": "Office_Order", "circular": "Circular", "policy": "Policy"}
    fname = f"{fname_map[doc.document_type]}_{doc.serial_number}.docx"
    return make_docx_response(word_doc, fname)
