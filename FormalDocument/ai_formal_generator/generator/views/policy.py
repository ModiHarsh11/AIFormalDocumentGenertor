"""Views for Policy document generation, preview, and export."""

import base64
import logging
import os
import uuid
from io import BytesIO

from django.conf import settings
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.template.loader import render_to_string
from django.utils import timezone
from django.views.decorators.http import require_POST
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from pypdf import PdfReader, PdfWriter
from weasyprint import HTML

from generator.constants import DESIGNATION_MAP

from .common import (
    POLICY, format_date_ddmmmyyyy,
    handle_generate_body, handle_regenerate_body, handle_update_body,
    maybe_translate_subject,
    get_logo_path, get_logo_file_uri, make_docx_response,
    save_document_to_db,
)

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# AI body generation
# ---------------------------------------------------------------------------

@require_POST
@login_required
def generate_policy_body(request):
    """Generate a Policy body via the AI service."""
    return handle_generate_body(request, "policy", "generate_policy_body")


@require_POST
@login_required
def regenerate_policy_body(request):
    """Regenerate a Policy body with refinement feedback."""
    return handle_regenerate_body(request, "policy", "regenerate_policy_body")


# ---------------------------------------------------------------------------
# Form & preview
# ---------------------------------------------------------------------------

@login_required
def result_policy(request):
    if request.method != "POST":
        return redirect("home")

    lang = request.POST.get("language") or "en"
    raw_date = request.POST.get("date") or timezone.now().strftime("%Y-%m-%d")
    date = format_date_ddmmmyyyy(raw_date) if raw_date else timezone.now().strftime("%d-%b-%Y").upper()
    subject = maybe_translate_subject(request.POST.get("subject"), lang)
    body = request.POST.get("body")
    body_prompt = request.POST.get("body_prompt", "").strip()
    attached_pdf_name = request.POST.get("attached_pdf_name", "")

    from_position = request.POST.get("from_position")
    from_designation = DESIGNATION_MAP.get(from_position, {}).get(lang, "") if from_position else ""

    to_recipients = request.POST.getlist("to_recipients[]")
    to_designations = [DESIGNATION_MAP.get(x, {}).get(lang, x) for x in to_recipients]

    lang_key = "hi" if lang == "hi" else "en"
    header = {
        "org_name": POLICY["header"][lang_key][0],
        "ministry": POLICY["header"][lang_key][1],
        "government": POLICY["header"][lang_key][2],
    }

    # Handle PDF upload
    uploaded_pdf = request.FILES.get("policy_pdf")
    pdf_path = None
    pdf_base64 = None
    if uploaded_pdf:
        upload_dir = os.path.join(settings.MEDIA_ROOT, "policy_uploads")
        os.makedirs(upload_dir, exist_ok=True)
        pdf_filename = f"policy_{uuid.uuid4().hex}.pdf"
        pdf_path = os.path.join(upload_dir, pdf_filename)
        with open(pdf_path, "wb+") as destination:
            for chunk in uploaded_pdf.chunks():
                destination.write(chunk)
        with open(pdf_path, "rb") as f:
            pdf_base64 = base64.b64encode(f.read()).decode("utf-8")

    data = {
        "language": lang,
        "header": header,
        "date": date,
        "date_raw": raw_date,
        "subject": subject,
        "body": body,
        "body_prompt": body_prompt,
        "from": from_designation,
        "from_position": from_position or "",
        "to_designations": to_designations,
        "to_data": to_recipients,
        "attached_pdf_name": attached_pdf_name,
        "uploaded_pdf_path": pdf_path,
    }

    # Persist to database
    doc = save_document_to_db(
        user=request.user,
        document_type="policy",
        language=lang,
        date_raw=raw_date,
        subject=subject,
        body_prompt=body_prompt,
        body=body or "",
        from_position=from_position or "",
        to_data=to_recipients,
        extra_data={
            "attached_pdf_name": attached_pdf_name,
            "uploaded_pdf_path": pdf_path or "",
        },
    )
    data["serial_number"] = doc.serial_number
    data["doc_id"] = doc.id

    # Store in session (without large base64)
    request.session["policy_data"] = data

    # Add pdf_base64 only for template rendering (not in session)
    if pdf_base64:
        data["pdf_base64"] = pdf_base64

    return render(request, "generator/result_policy.html", data)


@require_POST
@login_required
def update_policy_body(request):
    return handle_update_body(request, "policy_data")


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------

def download_policy_pdf(request):
    data = request.session.get("policy_data")
    if not data:
        return HttpResponse("No policy generated", status=400)

    # Work on a copy to avoid polluting the session with logo_path
    render_data = {**data}
    logo_uri = get_logo_file_uri()
    if logo_uri:
        render_data["logo_path"] = logo_uri

    html = render_to_string("generator/pdf_policy.html", render_data)

    first_page_pdf = HTML(
        string=html,
        base_url=str(settings.BASE_DIR),
    ).write_pdf(
        optimize_images=True,
        jpeg_quality=85,
        presentational_hints=True,
    )

    uploaded_pdf_path = data.get("uploaded_pdf_path")

    if uploaded_pdf_path and os.path.exists(uploaded_pdf_path):
        try:
            pdf_writer = PdfWriter()

            first_page_reader = PdfReader(BytesIO(first_page_pdf))
            for page in first_page_reader.pages:
                pdf_writer.add_page(page)

            uploaded_reader = PdfReader(uploaded_pdf_path)
            for page in uploaded_reader.pages:
                pdf_writer.add_page(page)

            output = BytesIO()
            pdf_writer.write(output)
            output.seek(0)
            final_pdf = output.read()
            # NOTE: do NOT delete the file here — DOCX export also needs it
        except Exception as e:
            logger.error("Error merging PDFs: %s", e)
            final_pdf = first_page_pdf
    else:
        final_pdf = first_page_pdf

    response = HttpResponse(final_pdf, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="Policy.pdf"'
    return response


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

def download_policy_docx(request):
    data = request.session.get("policy_data")
    if not data:
        return HttpResponse("No policy generated", status=400)

    doc = Document()
    lang = data.get("language", "en")

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Logo
    logo_path = get_logo_path()
    if logo_path:
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, height=Inches(0.9))
        doc.add_paragraph()

    # Header
    for line in data["header"].values():
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(14)

    # Title
    title_text = "नीति" if lang == "hi" else "Policy"
    p = doc.add_paragraph(title_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].underline = True
    p.runs[0].font.size = Pt(16)

    # Date
    date_label = "दिनांक :" if lang == "hi" else "Date :"
    p = doc.add_paragraph(f"{date_label} {data['date']}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    # Subject
    subject_label = "विषय :" if lang == "hi" else "Subject :"
    p = doc.add_paragraph(f"{subject_label} {data['subject']}")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    # Body
    p = doc.add_paragraph(data["body"])
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].font.size = Pt(12)

    # From
    p = doc.add_paragraph(data["from"])
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    doc.add_paragraph()

    # To
    to_label = "प्रति :" if lang == "hi" else "To :"
    p = doc.add_paragraph(to_label)
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    for designation in data.get("to_designations", []):
        p = doc.add_paragraph(designation, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    # Attached
    attached_label = "संलग्न :" if lang == "hi" else "Attached :"
    p = doc.add_paragraph(f"{attached_label} {data['attached_pdf_name']}")
    p.runs[0].font.size = Pt(12)

    # Try to convert and add PDF pages as images
    uploaded_pdf_path = data.get("uploaded_pdf_path")
    if uploaded_pdf_path and os.path.exists(uploaded_pdf_path):

        def _add_pdf_fallback_placeholder(message: str | None = None):
            """Add a placeholder page when PDF embedding fails."""
            doc.add_page_break()
            p = doc.add_paragraph("── Attached PDF ──")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True

            p = doc.add_paragraph(
                f"File: {data.get('attached_pdf_name', 'Attached Document')}"
            )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            note = message or "Please download the PDF version to view the complete document with all attached pages."
            p = doc.add_paragraph(note)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(10)

        try:
            import fitz  # pymupdf

            temp_dir = os.path.join(settings.MEDIA_ROOT, "policy_uploads")
            os.makedirs(temp_dir, exist_ok=True)

            doc.add_page_break()

            try:
                pdf_doc = fitz.open(uploaded_pdf_path)

                for idx in range(len(pdf_doc)):
                    page = pdf_doc[idx]
                    mat = fitz.Matrix(2.0, 2.0)  # 2x zoom ~ 144 dpi
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    temp_img_path = os.path.join(
                        temp_dir, f"temp_page_{uuid.uuid4().hex}_{idx}.png"
                    )
                    try:
                        pix.save(temp_img_path)
                        if idx > 0:
                            doc.add_page_break()
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(temp_img_path, width=Inches(6.5))
                        try:
                            os.remove(temp_img_path)
                        except Exception as rm_error:
                            logger.warning("Cleanup error: %s", rm_error)
                    except Exception as img_error:
                        logger.error("Image processing error for page %d: %s", idx, img_error)
                        continue

                pdf_doc.close()

                # Clean up uploaded PDF and clear session path after successful DOCX processing
                try:
                    os.remove(uploaded_pdf_path)
                    if "policy_data" in request.session:
                        request.session["policy_data"]["uploaded_pdf_path"] = None
                        request.session.modified = True
                except Exception as cleanup_error:
                    logger.warning("PDF cleanup error: %s", cleanup_error)

            except Exception as conv_error:
                logger.error("PDF conversion error: %s", conv_error)
                _add_pdf_fallback_placeholder(
                    "⚠ PDF pages could not be embedded as images. "
                    "Please download the PDF version for the complete document."
                )

        except ImportError as imp_err:
            logger.error("pymupdf not available: %s", imp_err)
            _add_pdf_fallback_placeholder()

        except Exception as e:
            logger.error("Error processing PDF attachment: %s", e)
            _add_pdf_fallback_placeholder(
                "[Error loading attached PDF — download as PDF for complete document]"
            )

    return make_docx_response(doc, "Policy.docx")

