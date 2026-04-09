"""Views for Advertisement document generation, preview, and export."""

import logging

from django.conf import settings
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.template.loader import render_to_string
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches, Twips
from weasyprint import HTML

from .common import (
    format_date_ddmmyyyy,
    get_logo_path, get_logo_file_uri, make_docx_response,
    save_document_to_db,
)

logger = logging.getLogger(__name__)


def _load_advertisement_data():
    """Load advertisement configuration data."""
    import json
    from pathlib import Path
    data_path = Path(settings.BASE_DIR) / "generator" / "data" / "advertisement.json"
    with data_path.open(encoding="utf-8") as f:
        return json.load(f)


def _format_ctc(ctc_value: int, ctc_type: str, lang: str = "en") -> str:
    """Format CTC value for display."""
    if ctc_type == "monthly":
        if lang == "hi":
            return f"₹{ctc_value:,} प्रति माह"
        return f"₹{ctc_value:,} Per Month"
    else:
        if lang == "hi":
            return f"₹{ctc_value:,} प्रति वर्ष"
        return f"₹{ctc_value:,} Per Year"


def _get_designation_label(designation_value: str, designations: list, lang: str = "en") -> str:
    """Get the display label for a designation value."""
    for d in designations:
        if d["value"] == designation_value:
            return d.get(f"label_{lang}", d.get("label_en", designation_value))
    return designation_value


def _get_qualification_labels(qual_values: list, qualifications: list, lang: str = "en") -> list:
    """Get display labels for selected qualifications."""
    labels = []
    for qv in qual_values:
        for q in qualifications:
            if q["value"] == qv:
                labels.append(q.get(f"label_{lang}", q.get("label_en", qv)))
                break
    return labels


@login_required
def result_advertisement(request):
    """Process advertisement form and render preview."""
    if request.method != "POST":
        return redirect("home")

    ad_data = _load_advertisement_data()
    lang = request.POST.get("language", "en")
    
    # Extract form data
    designation = request.POST.get("designation", "")
    num_requirements = request.POST.get("num_requirements", "1")
    qualifications = request.POST.getlist("qualifications[]")
    ctc_type = request.POST.get("ctc_type", "monthly")
    ctc_value = request.POST.get("ctc_value", "")
    years_experience = request.POST.get("years_experience", "0")
    govt_benefit = request.POST.get("govt_benefit", "no")
    email_id = request.POST.get("email_id", "")
    last_date = request.POST.get("last_date", "")
    
    # Convert and validate
    try:
        ctc_value_int = int(ctc_value) if ctc_value else 0
        num_req_int = int(num_requirements) if num_requirements else 1
        years_exp_int = int(years_experience) if years_experience else 0
    except ValueError:
        ctc_value_int = 0
        num_req_int = 1
        years_exp_int = 0

    # Format date
    formatted_date = format_date_ddmmyyyy(last_date) if last_date else ""

    # Get display labels
    designation_label = _get_designation_label(designation, ad_data["designations"], lang)
    qualification_labels = _get_qualification_labels(qualifications, ad_data["qualifications"], lang)
    ctc_display = _format_ctc(ctc_value_int, ctc_type, lang)

    # Get conditional text based on govt benefit selection
    lang_key = lang if lang in ["en", "hi"] else "en"
    govt_benefit_text = ad_data["govt_benefit_text"]["yes" if govt_benefit == "yes" else "no"][lang_key]
    application_text = ad_data["application_text"][lang_key]
    subject_line_note = ad_data["subject_line_note"][lang_key]
    last_date_text = ad_data["last_date_text"][lang_key].replace("{date}", formatted_date)

    # Prepare header
    header = ad_data["header"][lang_key]
    title = ad_data["title"][lang_key]

    data = {
        "language": lang,
        "header": header,
        "title": title,
        "designation": designation,
        "designation_label": designation_label,
        "num_requirements": num_req_int,
        "qualifications": qualifications,
        "qualification_labels": qualification_labels,
        "qualification_display": ", ".join(qualification_labels),
        "ctc_type": ctc_type,
        "ctc_value": ctc_value_int,
        "ctc_display": ctc_display,
        "years_experience": years_exp_int,
        "govt_benefit": govt_benefit,
        "govt_benefit_text": govt_benefit_text,
        "email_id": email_id,
        "last_date": formatted_date,
        "last_date_raw": last_date,
        "application_text": application_text,
        "subject_line_note": subject_line_note,
        "last_date_text": last_date_text,
    }

    # Save to database
    doc = save_document_to_db(
        user=request.user,
        document_type="advertisement",
        language=lang,
        date_raw=last_date,
        subject=f"Hiring Advertisement - {designation_label}",
        body=f"Hiring {num_req_int} {designation_label}",
        extra_data={
            "designation": designation,
            "num_requirements": num_req_int,
            "qualifications": qualifications,
            "ctc_type": ctc_type,
            "ctc_value": ctc_value_int,
            "years_experience": years_exp_int,
            "govt_benefit": govt_benefit,
            "email_id": email_id,
        },
    )
    data["serial_number"] = doc.serial_number
    data["doc_id"] = doc.id

    request.session["advertisement_data"] = data
    return render(request, "generator/result_advertisement.html", data)


def download_advertisement_pdf(request):
    """Generate and download PDF version of the advertisement."""
    data = request.session.get("advertisement_data")
    if not data:
        return HttpResponse("No advertisement generated", status=400)

    render_data = {**data}
    logo_uri = get_logo_file_uri()
    if logo_uri:
        render_data["logo_path"] = logo_uri

    html = render_to_string("generator/pdf_advertisement.html", render_data)

    pdf = HTML(
        string=html,
        base_url=str(settings.BASE_DIR),
    ).write_pdf(
        optimize_images=True,
        jpeg_quality=85,
        presentational_hints=True,
    )

    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="Advertisement.pdf"'
    return response


def download_advertisement_docx(request):
    """Generate and download DOCX version of the advertisement."""
    data = request.session.get("advertisement_data")
    if not data:
        return HttpResponse("No advertisement generated", status=400)

    doc = Document()
    lang = data.get("language", "en")

    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add BISAG Logo (centered)
    logo_path = get_logo_path()
    if logo_path:
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, height=Inches(0.9))

    # Header lines (Bold + Center)
    header = data["header"]
    for key in ["org_name", "ministry", "government"]:
        p = doc.add_paragraph(header[key])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)

    # Add spacing
    doc.add_paragraph()

    # Title (Bold + Center)
    title_p = doc.add_paragraph(data["title"])
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.runs[0].bold = True
    title_p.runs[0].underline = True
    title_p.runs[0].font.size = Pt(14)

    # Add spacing
    doc.add_paragraph()

    # Create Table
    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Table headers
    headers = [
        ("Designation", "पद"),
        ("No. of Requirements", "आवश्यकताओं की संख्या"),
        ("Educational Qualification", "शैक्षिक योग्यता"),
        ("Years of Experience", "अनुभव (वर्ष)"),
        (f"CTC ({data['ctc_type'].title()})", f"वेतन ({data['ctc_type'].title()})")
    ]
    
    hdr_cells = table.rows[0].cells
    for i, (en_text, hi_text) in enumerate(headers):
        text = hi_text if lang == "hi" else en_text
        hdr_cells[i].text = text
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Table data row
    row_cells = table.rows[1].cells
    row_cells[0].text = data["designation_label"]
    row_cells[1].text = str(data["num_requirements"])
    row_cells[2].text = data["qualification_display"]
    row_cells[3].text = str(data["years_experience"])
    row_cells[4].text = data["ctc_display"]

    # Center align all data cells
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(0.8)
        row.cells[2].width = Inches(2.5)
        row.cells[3].width = Inches(0.8)
        row.cells[4].width = Inches(1.2)

    # Add spacing
    doc.add_paragraph()

    # Government Benefit Text
    govt_p = doc.add_paragraph(data["govt_benefit_text"])
    govt_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    govt_p.runs[0].font.size = Pt(11)

    # Add spacing
    doc.add_paragraph()

    # Application Instructions
    app_p = doc.add_paragraph(data["application_text"])
    app_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    app_p.runs[0].font.size = Pt(11)

    # Add spacing
    doc.add_paragraph()

    # Email ID (Bold)
    email_label = "Email-id: " if lang == "en" else "ईमेल-आईडी: "
    email_p = doc.add_paragraph()
    label_run = email_p.add_run(email_label)
    label_run.font.size = Pt(11)
    email_run = email_p.add_run(data["email_id"])
    email_run.bold = True
    email_run.font.size = Pt(11)

    # Subject line note
    note_p = doc.add_paragraph(data["subject_line_note"])
    note_p.runs[0].italic = True
    note_p.runs[0].font.size = Pt(10)

    # Add spacing
    doc.add_paragraph()

    # Last date text
    last_p = doc.add_paragraph(data["last_date_text"])
    last_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    last_p.runs[0].font.size = Pt(11)

    return make_docx_response(doc, "Advertisement.docx")
